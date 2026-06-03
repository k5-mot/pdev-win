#!/usr/bin/env python3
"""Convert a DOCX file to Markdown and PNG assets."""

from __future__ import annotations

import argparse
import io
import json
import re
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from PIL import Image


def md_escape_cell(text: str) -> str:
    return re.sub(r"\s+", " ", text).replace("|", r"\|").strip()


def save_png(blob: bytes, images_dir: Path, index: int) -> str:
    images_dir.mkdir(parents=True, exist_ok=True)
    out = images_dir / f"image-{index:04d}.png"
    with Image.open(io.BytesIO(blob)) as image:
        if image.mode not in ("RGB", "RGBA"):
            image = image.convert("RGBA")
        image.save(out, "PNG")
    return f"images/{out.name}"


def paragraph_to_markdown(paragraph, rels, images_dir: Path, image_index: int):
    pieces = []
    warnings = []
    for run in paragraph.runs:
        for blip in run._element.xpath(".//a:blip"):
            rel_id = blip.get(qn("r:embed"))
            if rel_id and rel_id in rels:
                try:
                    image_index += 1
                    rel_path = save_png(rels[rel_id].blob, images_dir, image_index)
                    pieces.append(f"![image {image_index}]({rel_path})")
                except Exception as exc:
                    warnings.append(f"Could not export DOCX image {rel_id}: {exc}")
        if run.text:
            text = run.text
            if run.bold:
                text = f"**{text}**"
            if run.italic:
                text = f"*{text}*"
            pieces.append(text)

    text = "".join(pieces).strip()
    if not text:
        return "", image_index, warnings

    style = paragraph.style.name if paragraph.style else ""
    match = re.match(r"Heading\s+([1-6])", style or "")
    if match:
        text = f"{'#' * int(match.group(1))} {text}"
    elif style and "List Bullet" in style:
        text = f"- {text}"
    elif style and "List Number" in style:
        text = f"1. {text}"
    return text, image_index, warnings


def table_to_markdown(table) -> str:
    rows = []
    for row in table.rows:
        rows.append([md_escape_cell(cell.text) for cell in row.cells])
    if not rows:
        return ""
    width = max(len(row) for row in rows)
    rows = [row + [""] * (width - len(row)) for row in rows]
    header = "| " + " | ".join(rows[0]) + " |"
    sep = "| " + " | ".join(["---"] * width) + " |"
    body = ["| " + " | ".join(row) + " |" for row in rows[1:]]
    return "\n".join([header, sep] + body)


def iter_blocks(document):
    body = document.element.body
    para_map = {p._element: p for p in document.paragraphs}
    table_map = {t._element: t for t in document.tables}
    for child in body.iterchildren():
        if child in para_map:
            yield "paragraph", para_map[child]
        elif child in table_map:
            yield "table", table_map[child]


def convert(input_path: Path, output_dir: Path) -> dict:
    output_dir.mkdir(parents=True, exist_ok=True)
    images_dir = output_dir / "images"
    document = Document(input_path)
    rels = document.part.related_parts
    markdown = []
    warnings = []
    image_index = 0

    for kind, block in iter_blocks(document):
        if kind == "paragraph":
            text, image_index, block_warnings = paragraph_to_markdown(block, rels, images_dir, image_index)
            warnings.extend(block_warnings)
            if text:
                markdown.append(text)
        elif kind == "table":
            text = table_to_markdown(block)
            if text:
                markdown.append(text)

    md_path = output_dir / "document.md"
    md_path.write_text("\n\n".join(markdown) + "\n", encoding="utf-8")
    manifest = {
        "source": str(input_path),
        "markdown": str(md_path),
        "image_count": image_index,
        "warnings": warnings,
    }
    (output_dir / "manifest.json").write_text(json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8")
    return manifest


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("input", type=Path)
    parser.add_argument("output_dir", type=Path)
    args = parser.parse_args()
    manifest = convert(args.input, args.output_dir)
    print(json.dumps(manifest, ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
